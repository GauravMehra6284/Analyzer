from django.shortcuts import render
from django.http import JsonResponse
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from rest_framework import status, viewsets, filters
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend

from .models import Resume, ResumeAnalysis, Skill
from .serializers import ResumeAnalysisSerializer, SkillGapSerializer
from .resume_analysis import extract_text_from_pdf, analyze_resume_with_llm


# ✅ Test route to verify API health
def sample_data(request):
    return JsonResponse({'message': 'Hello from Django Gaurav!'})


# ✅ Upload a resume (PDF), analyze it with LLM, and save analysis
class ResumeUploadAPIView(APIView):
    parser_classes = [MultiPartParser]

    def post(self, request):
        uploaded_file = request.FILES.get("file")

        # ❌ Invalid or missing file
        if not uploaded_file or not uploaded_file.name.endswith(".pdf"):
            return Response({"error": "Only PDF files allowed."}, status=400)

        try:
            # ✅ Save file to Resume model (optional file storage)
            resume = Resume(file=uploaded_file)
            resume.save()

            # ✅ Extract text from the PDF
            path = resume.file.path
            extracted_text = extract_text_from_pdf(path)

            if not extracted_text.strip():
                return Response({"error": "No readable text found in resume."}, status=400)

            # ✅ Analyze resume text using Hugging Face LLM
            result = analyze_resume_with_llm(extracted_text)
            print("DEBUG ANALYSIS RESULT:", result)

            # ❌ If LLM fails
            if "error" in result:
                return Response({"error": result["error"]}, status=500)

            # ✅ Save results to ResumeAnalysis model
            analysis = ResumeAnalysis.objects.create(
                file_name=uploaded_file.name,
                ats_score=result.get("ats_score"),
                clarity_score=result.get("clarity_score"),
                experience=result.get("experience"),
                education=result.get("education"),
                skills_match=result.get("skills"),
                strengths=result.get("strengths"),
                weaknesses=result.get("weaknesses"),
                job_matches=result.get("job_matches"),
                status="completed",
                trend="neutral",
                previous_score=0
            )

            serializer = ResumeAnalysisSerializer(analysis)
            return Response({
                "message": "Resume uploaded and analyzed successfully.",
                "analysis": serializer.data
            }, status=201)

        except Exception as e:
            print("ERROR:", str(e))
            return Response({"error": "Something went wrong during analysis."}, status=500)


# ✅ View all resume analysis records with filters
class ResumeAnalysisViewSet(viewsets.ModelViewSet):
    queryset = ResumeAnalysis.objects.all().order_by('-upload_date')
    serializer_class = ResumeAnalysisSerializer
    filter_backends = [filters.SearchFilter, DjangoFilterBackend]
    search_fields = ['file_name']
    filterset_fields = ['status']


# ✅ Secure skill gap analysis for logged-in users
class SkillGapAnalysisView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        skills = Skill.objects.prefetch_related('courses').all()
        serializer = SkillGapSerializer(skills, many=True, context={'request': request})
        return Response(serializer.data)


# ✅ Render React frontend (index.html)
def react_index(request):
    return render(request, "index.html")


from rest_framework.decorators import api_view
from rest_framework.response import Response
from .models import ResumeAnalysis
from .serializers import ResumeAnalysisSerializer

@api_view(['GET'])
def latest_resume_analysis(request):
    latest = ResumeAnalysis.objects.order_by('-upload_date').first()
    if latest:
        serializer = ResumeAnalysisSerializer(latest)
        return Response(serializer.data)
    return Response({"error": "No resume analysis found"}, status=404)


# views.py
from rest_framework.decorators import api_view
from rest_framework.response import Response
from .models import ResumeAnalysis
from .serializers import ResumeAnalysisSerializer

@api_view(['GET'])
def get_latest_analysis(request):
    try:
        analysis = ResumeAnalysis.objects.latest('upload_date')  # You can filter by user later
        serializer = ResumeAnalysisSerializer(analysis)
        return Response({"analysis": serializer.data})
    except ResumeAnalysis.DoesNotExist:
        return Response({"analysis": None})

from rest_framework.decorators import api_view
from rest_framework.response import Response
from .models import ResumeAnalysis
from django.db.models import Avg
from datetime import datetime


@api_view(['GET'])
def sample_api(request):
    return Response({"message": "Server running successfully!"})


@api_view(['GET'])
def current_analysis(request):
    latest = ResumeAnalysis.objects.order_by('-upload_date').first()

    if not latest:
        return Response({"message": "No analysis found"}, status=204)

    return Response({
        "fileName": latest.file_name,
        "atsScore": latest.ats_score,
        "clarityScore": latest.clarity_score,
        "skills": latest.skills_match or [],
        "strengths": latest.strengths or [],
        "weaknesses": latest.weaknesses or [],
        "education": latest.education or "Not provided",
        "experience": latest.experience or "Not provided",
        "trend": latest.trend or "neutral"
    })


@api_view(['GET'])
def dashboard_stats(request):
    total = ResumeAnalysis.objects.count()
    avg_ats = ResumeAnalysis.objects.aggregate(Avg('ats_score'))['ats_score__avg'] or 0

    stats = [
        {
            "name": "Total Resumes Analyzed",
            "value": total,
            "change": "12%",
            "color": "bg-blue-500",
            "icon": "DocumentTextIcon"
        },
        {
            "name": "Average ATS Score",
            "value": f"{avg_ats:.0f}%",
            "change": "6%",
            "color": "bg-green-500",
            "icon": "ChartBarIcon"
        },
        {
            "name": "Job Matches Found",
            "value": total * 2,  # Placeholder logic
            "change": "18%",
            "color": "bg-purple-500",
            "icon": "BriefcaseIcon"
        },
        {
            "name": "Skills Improved",
            "value": total * 3,  # Placeholder logic
            "change": "9%",
            "color": "bg-orange-500",
            "icon": "AcademicCapIcon"
        }
    ]

    return Response(stats)


@api_view(['GET'])
def recent_analyses(request):
    recent = ResumeAnalysis.objects.order_by('-upload_date')[:5]
    data = []

    for analysis in recent:
        data.append({
            "id": analysis.id,
            "name": analysis.file_name,
            "score": int(analysis.ats_score or 0),
            "status": analysis.status.title(),
            "date": analysis.upload_date.strftime("%b %d, %Y %I:%M %p")
        })

    return Response(data)




# views.py
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from docx import Document
from PyPDF2 import PdfReader
import os

@api_view(['POST'])
@parser_classes([MultiPartParser])
def upload_jd_file(request):
    file = request.FILES.get('file')
    if not file:
        return Response({"error": "No file uploaded"}, status=400)

    try:
        filename = file.name.lower()

        if filename.endswith('.pdf'):
            reader = PdfReader(file)
            text = "\n".join([page.extract_text() or "" for page in reader.pages])

        elif filename.endswith('.docx'):
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])

        elif filename.endswith('.txt'):
            text = file.read().decode("utf-8")

        else:
            return Response({"error": "Unsupported file format. Only PDF, DOCX, TXT supported."}, status=400)

        return Response({"jd_text": text})

    except Exception as e:
        return Response({"error": str(e)}, status=500)


from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from PyPDF2 import PdfReader
from docx import Document
import os
import requests


import os
import json
import requests
from docx import Document
from PyPDF2 import PdfReader

from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response


from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from docx import Document
from PyPDF2 import PdfReader
import requests
import os
import json

@api_view(['POST'])
@parser_classes([MultiPartParser])
def match_resume_jd(request):
    resume_file = request.FILES.get('resume')
    jd_file = request.FILES.get('jd')

    if not resume_file or not jd_file:
        return Response({"error": "Both resume and JD files are required"}, status=400)

    try:
        # --- Extract Resume Text ---
        if resume_file.name.endswith('.pdf'):
            resume_text = "\n".join([page.extract_text() or "" for page in PdfReader(resume_file).pages])
        elif resume_file.name.endswith('.docx'):
            doc = Document(resume_file)
            resume_text = "\n".join([p.text for p in doc.paragraphs])
        elif resume_file.name.endswith('.txt'):
            resume_text = resume_file.read().decode('utf-8')
        else:
            return Response({"error": "Unsupported resume file format"}, status=400)

        # --- Extract JD Text ---
        if jd_file.name.endswith('.pdf'):
            jd_text = "\n".join([page.extract_text() or "" for page in PdfReader(jd_file).pages])
        elif jd_file.name.endswith('.docx'):
            doc = Document(jd_file)
            jd_text = "\n".join([p.text for p in doc.paragraphs])
        elif jd_file.name.endswith('.txt'):
            jd_text = jd_file.read().decode('utf-8')
        else:
            return Response({"error": "Unsupported JD file format"}, status=400)

        if not resume_text.strip() or not jd_text.strip():
            return Response({"error": "Extracted text is empty from resume or JD."}, status=400)

        # --- Prompt for LLM ---
        prompt = f"""
You are an AI assistant that matches resumes with job descriptions.

Here is the RESUME:
\"\"\"{resume_text}\"\"\"

Here is the JOB DESCRIPTION:
\"\"\"{jd_text}\"\"\"

Instructions:
- Compare the resume and job description.
- Give a match score from 0 to 100.
- Suggest 3 improvements to make the resume better match the job description.

Return ONLY a JSON object with this format (and nothing else):
{{
  "match_score": 85,
  "suggestions": [
    "Add experience with cloud technologies",
    "Include project management skills",
    "Highlight familiarity with Docker and Kubernetes"
  ]
}}
        """.strip()

        # --- OpenRouter API Call ---
        OPENROUTER_API_KEY = os.getenv("OR_API_KEY")
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }

        body = {
            "model": "meta-llama/llama-3-70b-instruct",
            "temperature": 0,  # 🔐 Deterministic output
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }

        response = requests.post("https://openrouter.ai/api/v1/chat/completions", json=body, headers=headers)
        response.raise_for_status()
        result = response.json()

        # --- Extract and Parse LLM Response ---
        content = result["choices"][0]["message"]["content"]

        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            return Response({
                "error": "LLM returned non-JSON output",
                "raw_response": content
            }, status=500)

        return Response(data)

    except Exception as e:
        return Response({"error": str(e)}, status=500)


from rest_framework import generics
from django.contrib.auth.models import User
from rest_framework.serializers import ModelSerializer

class RegisterSerializer(ModelSerializer):
    class Meta:
        model = User
        fields = ('username', 'password', 'email')
        extra_kwargs = {'password': {'write_only': True}}

    def create(self, validated_data):
        user = User.objects.create_user(
            username=validated_data['username'],
            password=validated_data['password'],
            email=validated_data.get('email', '')
        )
        return user

class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = RegisterSerializer


from django.contrib.auth.models import User
from rest_framework import generics, serializers
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

class RegisterSerializer(serializers.ModelSerializer):
    class Meta:
        model = User
        fields = ['username', 'password', 'email']
        extra_kwargs = {'password': {'write_only': True}}

    def create(self, validated_data):
        user = User.objects.create_user(
            username=validated_data['username'],
            password=validated_data['password'],
            email=validated_data.get('email', '')
        )
        return user

class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = RegisterSerializer

class ResumeAnalysisView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Your resume analysis logic goes here
        return Response({"message": "Resume analyzed successfully."})

